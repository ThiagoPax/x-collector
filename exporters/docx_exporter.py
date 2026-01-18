"""Exportador para DOCX usando python-docx com relatório diagnóstico."""
from __future__ import annotations
import os
import asyncio
from datetime import datetime
from pathlib import Path
from typing import Optional, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from core.models import Post, CollectionResult, CollectionParams


class DocxExporter:
    """Exporta posts para documento Word com relatório diagnóstico."""
    
    def __init__(self, output_dir: str = None):
        """
        Inicializa o exportador.
        
        Args:
            output_dir: Diretório de saída (padrão: ./exports)
        """
        self.output_dir = Path(output_dir or os.getenv("EXPORTS_DIR", "./exports"))
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def export(
        self,
        result: CollectionResult,
        filename: str = None,
        include_diagnostic: bool = False,
        diagnostic_report = None,
    ) -> str:
        """
        Exporta resultado da coleta para DOCX.
        
        Args:
            result: Resultado da coleta
            filename: Nome do arquivo (opcional)
            include_diagnostic: Se deve incluir relatório diagnóstico (padrão: False - apenas posts)
            diagnostic_report: Relatório pré-gerado (opcional)
            
        Returns:
            Caminho do arquivo gerado
        """
        doc = Document()
        
        # Configurar estilos
        self._setup_styles(doc)
        
        # Cabeçalho
        self._add_header(doc, result)
        
        # Sumário com métricas completas
        self._add_summary(doc, result)
        
        # Relatório diagnóstico (se disponível)
        if include_diagnostic:
            self._add_diagnostic_report(doc, result, diagnostic_report)
        
        # Posts
        self._add_posts(doc, result.posts)
        
        # Gerar nome do arquivo
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"coleta_x_{timestamp}.docx"
        
        if not filename.endswith(".docx"):
            filename += ".docx"
        
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        
        return str(filepath)
    
    def _setup_styles(self, doc: Document):
        """Configura estilos do documento."""
        # Título
        try:
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(20)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(29, 161, 242)  # Azul Twitter
        except:
            pass
        
        # Subtítulo
        try:
            subtitle_style = doc.styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_style.font.size = Pt(12)
            subtitle_style.font.color.rgb = RGBColor(101, 119, 134)
        except:
            pass
        
        # Seção
        try:
            section_style = doc.styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(20, 23, 26)
        except:
            pass
    
    def _add_header(self, doc: Document, result: CollectionResult):
        """Adiciona cabeçalho do documento."""
        # Título principal
        try:
            title = doc.add_paragraph("🐦 Coleta de Posts do X", style='CustomTitle')
        except:
            title = doc.add_paragraph("🐦 Coleta de Posts do X")
            title.runs[0].font.size = Pt(20)
            title.runs[0].font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data da coleta
        date_str = result.started_at.strftime("%d/%m/%Y às %H:%M")
        try:
            subtitle = doc.add_paragraph(f"Gerado em {date_str}", style='CustomSubtitle')
        except:
            subtitle = doc.add_paragraph(f"Gerado em {date_str}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Query/URL em destaque
        query_para = doc.add_paragraph()
        query_para.add_run("📌 Pesquisa: ").bold = True
        query_para.add_run(result.query_or_url)
        
        doc.add_paragraph()
    
    def _add_summary(self, doc: Document, result: CollectionResult):
        """Adiciona sumário com métricas completas incluindo views."""
        doc.add_heading("📊 Resumo da Coleta", level=1)
        
        # Calcular métricas
        total_likes = sum(p.metrics.likes or 0 for p in result.posts)
        total_reposts = sum(p.metrics.reposts or 0 for p in result.posts)
        total_replies = sum(p.metrics.replies or 0 for p in result.posts)
        total_views = sum(p.metrics.views or 0 for p in result.posts)
        
        # Tabela de métricas
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        
        metrics_data = [
            ("📝 Total de posts", f"{result.total_collected:,}".replace(",", ".")),
            ("❤️ Total de curtidas", f"{total_likes:,}".replace(",", ".")),
            ("🔁 Total de reposts", f"{total_reposts:,}".replace(",", ".")),
            ("💬 Total de respostas", f"{total_replies:,}".replace(",", ".")),
            ("👁️ Total de visualizações", f"{total_views:,}".replace(",", ".")),
            ("📈 Média de curtidas/post", f"{total_likes/max(result.total_collected,1):.1f}"),
            ("📊 Média de views/post", f"{total_views/max(result.total_collected,1):.1f}"),
        ]
        
        for i, (label, value) in enumerate(metrics_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            # Negrito na primeira coluna
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Período da coleta
        if result.posts:
            dates = [p.datetime for p in result.posts if p.datetime]
            if dates:
                oldest = min(dates).strftime("%d/%m/%Y")
                newest = max(dates).strftime("%d/%m/%Y")
                doc.add_paragraph(f"📅 Período coberto: {oldest} a {newest}")
        
        # Duração
        if result.finished_at:
            duration = (result.finished_at - result.started_at).total_seconds()
            doc.add_paragraph(f"⏱️ Duração da coleta: {duration:.1f} segundos")
        
        # Filtros
        params = result.params
        filters = []
        filters.append(f"Ordenação: {'Mais recentes' if params.search_type.value == 'latest' else 'Mais relevantes'}")
        if params.max_posts:
            filters.append(f"Limite: {params.max_posts} posts")
        if params.max_days:
            filters.append(f"Período: últimos {params.max_days} dias")
        if params.language:
            filters.append(f"Idioma: {params.language}")
        
        doc.add_paragraph()
        filter_para = doc.add_paragraph()
        filter_para.add_run("⚙️ Filtros aplicados: ").bold = True
        filter_para.add_run(" | ".join(filters))
        
        doc.add_paragraph()
    
    def _add_diagnostic_report(self, doc: Document, result: CollectionResult, diagnostic_report = None):
        """Adiciona relatório diagnóstico ao documento."""
        # Gerar relatório se não fornecido
        if not diagnostic_report and result.posts:
            try:
                from core.analyzer import ContentAnalyzer
                analyzer = ContentAnalyzer()
                # Versão síncrona básica
                diagnostic_report = asyncio.run(
                    analyzer.analyze_posts(result.posts, result.query_or_url)
                )
            except Exception as e:
                print(f"⚠️ Não foi possível gerar relatório diagnóstico: {e}")
                return
        
        if not diagnostic_report:
            return
        
        # Título da seção
        doc.add_heading("📋 Relatório de Diagnóstico", level=1)
        
        # Valor percebido
        self._add_section(doc, "💎 Valor Percebido pelo Público Final", 
                         diagnostic_report.valor_percebido)
        
        # Mensagem principal
        self._add_section(doc, "📌 Mensagem Principal Identificada", 
                         diagnostic_report.mensagem_principal)
        
        # Submensagens
        if diagnostic_report.submensagens:
            self._add_list_section(doc, "📝 Submensagens Implícitas", 
                                  diagnostic_report.submensagens)
        
        # Possíveis vieses
        if diagnostic_report.possiveis_vieses:
            self._add_list_section(doc, "⚠️ Possíveis Vieses Identificados", 
                                  diagnostic_report.possiveis_vieses)
        
        # Pontos positivos
        if diagnostic_report.pontos_positivos:
            self._add_list_section(doc, "✅ Pontos Positivos", 
                                  diagnostic_report.pontos_positivos)
        
        # Pontos negativos
        if diagnostic_report.pontos_negativos:
            self._add_list_section(doc, "❌ Pontos Negativos / Limitações", 
                                  diagnostic_report.pontos_negativos)
        
        # Elementos de destaque
        if diagnostic_report.elementos_destaque:
            self._add_list_section(doc, "🌟 Elementos de Maior Destaque", 
                                  diagnostic_report.elementos_destaque)
        
        # Percepção de qualidade
        self._add_section(doc, "📊 Percepção Geral da Qualidade", 
                         diagnostic_report.percepcao_qualidade)
        
        # Observações
        if diagnostic_report.observacoes:
            self._add_list_section(doc, "💡 Observações para Tomada de Decisão", 
                                  diagnostic_report.observacoes)
        
        doc.add_paragraph()
        doc.add_paragraph("─" * 50)
        doc.add_paragraph()
    
    def _add_section(self, doc: Document, title: str, content: str):
        """Adiciona uma seção com título e conteúdo."""
        if not content:
            return
        para = doc.add_paragraph()
        para.add_run(title).bold = True
        doc.add_paragraph(content)
        doc.add_paragraph()
    
    def _add_list_section(self, doc: Document, title: str, items: List[str]):
        """Adiciona uma seção com lista."""
        if not items:
            return
        para = doc.add_paragraph()
        para.add_run(title).bold = True
        for item in items:
            doc.add_paragraph(f"• {item}")
        doc.add_paragraph()
    
    def _add_posts(self, doc: Document, posts: List[Post]):
        """Adiciona os posts ao documento."""
        doc.add_heading("📝 Posts Coletados", level=1)
        
        for i, post in enumerate(posts, 1):
            # Separador entre posts
            if i > 1:
                doc.add_paragraph("─" * 50)
            
            # Número e autor
            header = doc.add_paragraph()
            header.add_run(f"#{i} ").bold = True
            header.add_run(f"{post.author_name} ").bold = True
            header.add_run(f"(@{post.author_handle})")
            
            # Data
            if post.datetime:
                date_str = post.datetime.strftime("%d/%m/%Y às %H:%M")
                doc.add_paragraph(f"📅 {date_str}")
            
            # Texto do post
            if post.text:
                text_para = doc.add_paragraph()
                text_para.add_run(post.text)
            
            # Métricas (incluindo views)
            metrics = []
            if post.metrics.likes is not None:
                metrics.append(f"❤️ {post.metrics.likes:,}".replace(",", "."))
            if post.metrics.reposts is not None:
                metrics.append(f"🔁 {post.metrics.reposts:,}".replace(",", "."))
            if post.metrics.replies is not None:
                metrics.append(f"💬 {post.metrics.replies:,}".replace(",", "."))
            if post.metrics.views is not None:
                metrics.append(f"👁️ {post.metrics.views:,}".replace(",", "."))
            
            if metrics:
                metrics_para = doc.add_paragraph()
                metrics_para.add_run("📊 Engajamento: ").bold = True
                metrics_para.add_run(" | ".join(metrics))
            
            # Link do post
            link_para = doc.add_paragraph()
            link_para.add_run("🔗 ").bold = True
            self._add_hyperlink(link_para, post.url, post.url)
            
            # Hashtags
            if post.hashtags:
                doc.add_paragraph(f"#️⃣ {' '.join(post.hashtags)}")
            
            # Mentions
            if post.mentions:
                doc.add_paragraph(f"@ {' '.join(post.mentions)}")
            
            # Links externos
            if post.links:
                links_para = doc.add_paragraph()
                links_para.add_run("🔗 Links externos: ").bold = True
                for link in post.links[:3]:
                    doc.add_paragraph(f"  • {link}")
            
            # Mídia
            if post.media_urls:
                media_para = doc.add_paragraph()
                media_para.add_run("📷 Mídia: ").bold = True
                media_para.add_run(f"{len(post.media_urls)} arquivo(s)")
            
            # Flags (tipo de post)
            flags = []
            if post.is_repost:
                flags.append("🔁 Repost")
            if post.is_reply:
                flags.append("↩️ Resposta")
            if post.is_quote:
                flags.append("💬 Citação")
            
            if flags:
                doc.add_paragraph(" | ".join(flags))
    
    def _add_hyperlink(self, paragraph, url: str, text: str):
        """Adiciona hyperlink ao parágrafo."""
        run = paragraph.add_run(text)
        run.font.color.rgb = RGBColor(29, 161, 242)  # Azul Twitter
        run.font.underline = True


def export_to_docx(
    result: CollectionResult,
    output_dir: str = None,
    filename: str = None,
    include_diagnostic: bool = False,  # Por padrão, não incluir análise no arquivo
    diagnostic_report = None,
) -> str:
    """
    Função helper para exportar para DOCX.
    
    Args:
        result: Resultado da coleta
        output_dir: Diretório de saída
        filename: Nome do arquivo
        include_diagnostic: Se deve incluir relatório diagnóstico (padrão: False)
        diagnostic_report: Relatório pré-gerado (opcional)
        
    Returns:
        Caminho do arquivo gerado
    """
    exporter = DocxExporter(output_dir)
    return exporter.export(result, filename, include_diagnostic, diagnostic_report)
